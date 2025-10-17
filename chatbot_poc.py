# ======================= chatbot_poc.py (UPDATED) =======================
# Full resume analyzer CLI: parse files, extract skills (model + heuristics), ingest to Qdrant,
# semantic search, and JSON output for details & search.
# Put this file next to model.py and your 'resumes' folder and skills.csv.

import os
# Avoid TF being pulled in by transformers on Windows
os.environ["TRANSFORMERS_NO_TF"] = "1"
import re
import io 
import logging
from typing import List, Dict, Any, Optional
import json
import csv
import time
from difflib import get_close_matches
import numpy as _np
from collections import defaultdict
import uuid
import datetime 
import re as _re
from model import init_skill_model 

# ---------------- Optional deps (guarded imports) ----------------
try:
    import spacy 
except Exception as e:
    spacy = None
try:
    import fitz  # PyMuPDF
except Exception: 
    fitz = None
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from sentence_transformers import SentenceTransformer
except Exception as e:
    SentenceTransformer = None
    print("Warning: sentence-transformers not available:", e)

try:
    from qdrant_client import QdrantClient
    from qdrant_client.http import models as qm
except Exception as e:
    QdrantClient = None
    qm = None
    print("Warning: qdrant-client not available:", e)

# rapidfuzz optional for faster fuzzy matching
try:
    from rapidfuzz import process as rf_process, fuzz as rf_fuzz
    RAPIDFUZZ = True
except Exception:
    RAPIDFUZZ = False

# import model initializer (your local model.py)

# ---------------- CONFIG ----------------
HF_LOCAL_SNAPSHOT = None

# Default to a small, fast model; can override via EMBEDDING_MODEL env var
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "google/gemma-3-270m-it")

QDRANT_HOST = os.getenv("QDRANT_HOST", "localhost")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", 6333))
QDRANT_COLLECTION = os.getenv("QDRANT_COLLECTION", "resumes_collection")

RESUME_FOLDER = os.getenv("RESUME_FOLDER", "resumes")
SPACY_MODEL = os.getenv("SPACY_MODEL", "en_core_web_trf")
SKILLS_CSV_PATH = os.getenv("SKILLS_CSV_PATH", "categorized.csv")

# ---------------- NLP init ----------------
if spacy is not None:
    try:
        nlp = spacy.load(SPACY_MODEL)
        print(f"spaCy model '{SPACY_MODEL}' loaded.")
    except Exception as e:
        print("spaCy load failed; using blank model:", e)
        nlp = spacy.blank("en")
else:
    print("spaCy not installed; using minimal tokenizer.")
    class _Dummy:
        def __call__(self, x): return []
        @property
        def ents(self): return []
    nlp = _Dummy()

# ---------------- init skill model (force Gemma, transformers backend) ----------------
skill_pipe = None
if init_skill_model is not None:
    try:
        # Force the Gemma model and transformers backend explicitly to ensure 640-dim vectors
        skill_pipe = init_skill_model(
            model_id=EMBEDDING_MODEL,     # EMBEDDING_MODEL defaults to "google/gemma-3-270m-it"
            backend="transformers",       # force transformers/Gemma (not sentence-transformers)
            device=None,                  # None => auto select 'cuda' if available else 'cpu'
            normalize=True                # keep L2 normalization as implemented in model.py
        )
    except Exception as e:
        print("init_skill_model error (forcing Gemma):", e)
        skill_pipe = None

# show what we initialized (helpful debug)
if skill_pipe is not None:
    try:
        print("skill_pipe initialized. embedding dim:", skill_pipe.get_sentence_embedding_dimension())
    except Exception:
        print("skill_pipe initialized (dimension unknown).")
else:
    print("skill_pipe available: False")


# ---------------- Embeddings backend (SentenceTransformer OR Gemma) ----------------
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch

EMBEDDING_BACKEND = os.getenv("EMBEDDING_BACKEND", "sentence-transformers").lower()  # or "gemma"
GEMMA_LOCAL_PATH = os.getenv("GEMMA_LOCAL_PATH")  # optional local path

class _SentenceTransformerWrapper:
    def __init__(self, model_name):
        if SentenceTransformer is None:
            raise RuntimeError("sentence-transformers not installed")
        self._m = SentenceTransformer(model_name)
        self._dim = self._m.get_sentence_embedding_dimension()
    def encode(self, texts):
        single = isinstance(texts, str)
        vecs = self._m.encode([texts] if single else texts,
                              convert_to_numpy=True, normalize_embeddings=True)
        return vecs[0] if single else vecs
    def get_sentence_embedding_dimension(self):
        return int(self._dim)

class GemmaEmbedder:
    """
    Sentence embedding using a Gemma causal LM by mean-pooling the last hidden states.
    Works with any HF Gemma checkpoint (e.g., google/gemma-2-2b-it, gemma-2-9b-it, gemma-3-270m-it).
    """
    def __init__(self, model_id: str = None, device: str = None, dtype: str = "auto", max_length: int = 1024):
        self.model_id = model_id or os.getenv("GEMMA_MODEL_ID", "google/gemma-3-270m-it")
        self.max_length = max_length
        if device is None:
            device = "cuda" if torch.cuda.is_available() else "cpu"
        self.device = device

        # Load tokenizer + model
        self.tokenizer = AutoTokenizer.from_pretrained(self.model_id, use_fast=True)
        self.model = AutoModelForCausalLM.from_pretrained(
            self.model_id,
            torch_dtype="auto" if dtype == "auto" else getattr(torch, dtype),
            device_map=None,
        )
        self.model.eval()
        self.model.to(self.device)
        torch.set_grad_enabled(False)

        # Dimension = hidden size from config
        self.dim = getattr(self.model.config, "hidden_size", None)
        if not self.dim:
            tmp_inp = self.tokenizer("x", return_tensors="pt").to(self.device)
            out = self.model(**tmp_inp, output_hidden_states=True)
            self.dim = out.hidden_states[-1].shape[-1]

    def get_sentence_embedding_dimension(self):
        return int(self.dim)

    @torch.inference_mode()
    def encode(self, texts, batch_size: int = 16, normalize: bool = True):
        single = isinstance(texts, str)
        if single:
            texts = [texts]
        embs = []
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i+batch_size]
            tok = self.tokenizer(
                batch, max_length=self.max_length, padding=True, truncation=True, return_tensors="pt"
            ).to(self.device)
            outputs = self.model(**tok, output_hidden_states=True)
            last_hidden = outputs.hidden_states[-1]      # [B, T, H]
            mask = tok.attention_mask.unsqueeze(-1)      # [B, T, 1]
            summed = (last_hidden * mask).sum(dim=1)     # [B, H]
            counts = mask.sum(dim=1).clamp(min=1)        # [B, 1]
            pooled = summed / counts                     # [B, H]
            vecs = pooled.detach().float().cpu().numpy()
            if normalize:
                norms = _np.linalg.norm(vecs, axis=1, keepdims=True)
                norms = _np.where(norms == 0, 1.0, norms)
                vecs = vecs / norms
            embs.append(vecs)
        out = _np.vstack(embs)
        return out[0] if single else out  # [N, H]

# Choose backend
embedding_model = None
if EMBEDDING_BACKEND == "gemma":
    embedding_model = GemmaEmbedder(
        model_id=os.getenv("GEMMA_MODEL_ID", "google/gemma-3-270m-it"),
        max_length=int(os.getenv("GEMMA_MAX_LEN", "1024"))
    )
    EMBEDDING_DIM = embedding_model.get_sentence_embedding_dimension()
    def embed_texts(texts):
        return embedding_model.encode(texts)
else:
    # Fallback: MiniLM from sentence_transformers (and actually ASSIGN embedding_model)
    _minilm_id = os.getenv("MINILM_MODEL_ID", "sentence-transformers/all-MiniLM-L6-v2")
    embedding_model = _SentenceTransformerWrapper(_minilm_id)
    EMBEDDING_DIM = embedding_model.get_sentence_embedding_dimension()
    def embed_texts(texts):
        return embedding_model.encode(texts)

# Qdrant init (collection creation is handled by ingest_resume.py)
qdrant_client = None
if QdrantClient is not None and embedding_model is not None:
    try:
        qdrant_client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)
        print("Qdrant client initialized (collection creation handled by ingest_resume.py).")
    except Exception as e:
        print("Qdrant init failed:", e)
        qdrant_client = None

# --- CSV-based skill taxonomy (fallbacks + helpers) ----------------------
SKILL_TAXONOMY = {
    "programming_languages": ["python","java","c","c++","c#","javascript","typescript","go","scala","ruby","php","r","swift","kotlin"],
    "cloud": ["aws","amazon web services","azure","microsoft azure","gcp","google cloud","google cloud platform"],
    "databases": ["mysql","postgresql","postgres","oracle","sql server","mongodb","redis","dynamodb","bigquery","snowflake","redshift"],
    "data_engineering": ["hadoop","spark","pyspark","hive","airflow","databricks","dbt","kafka","sqoop"],
    "devops_ci_cd": ["docker","kubernetes","jenkins","gitlab ci","github actions","circleci","terraform","ansible","azure devops"],
    "web_frameworks": ["react","angular","vue","spring","django","flask","express","node","asp.net"],
    "analytics_and_viz": ["tableau","powerbi","looker","matplotlib","seaborn"],
    "testing_and_quality": ["pytest","junit","selenium","cucumber"],
    "storage_and_format": ["parquet","avro","orc","csv","json"],
    "other_tech": ["git","github","gitlab","bitbucket","rest","graphql","elasticsearch"],
    "soft_skills": ["communication","leadership","teamwork","collaboration","problem solving","adaptability","management","presentation","mentoring"],
}
CANONICAL_SKILLS: set = set()
SKILL_ALIASES: Dict[str, str] = {}
SKILL_TO_CATEGORY: Dict[str, str] = {}
CANONICAL_LIST: List[str] = []
CSV_SKILLS_LOADED = False
_FLAT_SKILL_TO_CATEGORY: Dict[str, str] = {}
_SOFT_SKILL_KEYWORDS: set = set()

def load_skills_from_csv(path: str = SKILLS_CSV_PATH) -> int:
    global CANONICAL_SKILLS, SKILL_ALIASES, SKILL_TO_CATEGORY, CANONICAL_LIST, CSV_SKILLS_LOADED
    CANONICAL_SKILLS = set(); SKILL_ALIASES = {}; SKILL_TO_CATEGORY = {}; CANONICAL_LIST = []
    CSV_SKILLS_LOADED = False

    if not path:
        print("load_skills_from_csv: no path provided")
        return 0

    if not os.path.exists(path):
        print(f"load_skills_from_csv: file not found at '{path}'")
        return 0

    def _cell_to_str(cell):
        """
        Safely convert a CSV cell to a stripped string.
        Handles None, list, numbers, etc.
        """
        if cell is None:
            return ""
        # If the CSV parser returned a list (weird), join it
        if isinstance(cell, (list, tuple)):
            try:
                cell = ", ".join(str(x) for x in cell)
            except Exception:
                cell = str(cell)
        try:
            return str(cell).strip()
        except Exception:
            return ""

    try:
        with open(path, newline='', encoding='utf-8', errors='replace') as f:
            sample = f.read(8192)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
                reader = csv.DictReader(f, dialect=dialect)
            except Exception:
                f.seek(0)
                reader = csv.DictReader(f)

            # header_map: lowercase->original
            header_map = {h.lower(): h for h in (reader.fieldnames or []) if h}

            # pick skill/category/alias columns robustly
            skill_col = header_map.get("skill") or header_map.get("skills") or header_map.get("name") or header_map.get("canonical") or ((reader.fieldnames or [None])[0] if reader.fieldnames else None)
            cat_col = header_map.get("category") or header_map.get("cat") or header_map.get("group") or header_map.get("type")
            alias_col = header_map.get("alias") or header_map.get("aliases") or header_map.get("synonyms") or header_map.get("alt")

            rows = 0
            for r in reader:
                rows += 1
                # Try to safely fetch values even if the row mapping is odd
                try:
                    raw_skill = _cell_to_str(r.get(skill_col) if skill_col in (r or {}) else None)
                except Exception:
                    # fallback: try to find first non-empty cell
                    raw_skill = ""
                    if isinstance(r, dict):
                        for v in r.values():
                            s = _cell_to_str(v)
                            if s:
                                raw_skill = s
                                break

                raw_cat = ""
                raw_alias = ""
                if isinstance(r, dict):
                    try:
                        raw_cat = _cell_to_str(r.get(cat_col)) if cat_col else ""
                    except Exception:
                        raw_cat = ""
                    try:
                        raw_alias = _cell_to_str(r.get(alias_col)) if alias_col else ""
                    except Exception:
                        raw_alias = ""

                if not raw_skill:
                    continue

                canon = raw_skill.lower()
                CANONICAL_SKILLS.add(canon)
                if raw_cat:
                    SKILL_TO_CATEGORY[canon] = raw_cat.lower()
                if raw_alias:
                    # split aliases by common delimiters and handle list-like strings
                    for a in _re.split(r'[;,/]+', raw_alias):
                        a = a.strip().lower()
                        if a:
                            SKILL_ALIASES[a] = canon

            CANONICAL_LIST = sorted(list(CANONICAL_SKILLS))
            CSV_SKILLS_LOADED = True
            print(f"Loaded {len(CANONICAL_SKILLS)} canonical skills from '{path}' and {len(SKILL_ALIASES)} aliases.")
            return len(CANONICAL_SKILLS)
    except Exception as e:
        print("Error reading skills CSV:", e)
        return 0

# load once (like before)
_loaded_count = load_skills_from_csv(SKILLS_CSV_PATH)


def _rebuild_flat_lookup_from_csv():
    global _FLAT_SKILL_TO_CATEGORY, _SOFT_SKILL_KEYWORDS
    _FLAT_SKILL_TO_CATEGORY = {}
    _SOFT_SKILL_KEYWORDS = set()

    if CSV_SKILLS_LOADED and CANONICAL_LIST:
        for canon in CANONICAL_LIST:
            cat = SKILL_TO_CATEGORY.get(canon, "") or ""
            key = canon.lower().strip()
            _FLAT_SKILL_TO_CATEGORY[key] = cat or "other_tech"

        for alias, canon in SKILL_ALIASES.items():
            key = alias.lower().strip()
            can_cat = SKILL_TO_CATEGORY.get(canon, "") or ""
            _FLAT_SKILL_TO_CATEGORY[key] = can_cat or "other_tech"

        for canon, cat in SKILL_TO_CATEGORY.items():
            if cat and ("soft" in cat or cat in ("soft_skills", "soft")):
                _SOFT_SKILL_KEYWORDS.add(canon.lower())
    else:
        for cat, toks in SKILL_TAXONOMY.items():
            if cat == "soft_skills":
                for s in toks: _SOFT_SKILL_KEYWORDS.add(s.lower())
            else:
                for t in toks: _FLAT_SKILL_TO_CATEGORY[t.lower()] = cat

_rebuild_flat_lookup_from_csv()
print(f"Flat lookup built. Flat keys: {len(_FLAT_SKILL_TO_CATEGORY)}; soft keywords: {len(_SOFT_SKILL_KEYWORDS)}")
if _loaded_count == 0:
    print("WARNING: No canonical skills loaded from categorized.csv. Check file contents/encoding and SKILLS_CSV_PATH.")
else:
    print("SUCCESS: canonical skills loaded.")

# ---------------- Canonical mapping helpers ----------------
_skill_emb_cache: Optional[Dict[str, _np.ndarray]] = None
def map_token_to_canonical(token: str, fuzzy_cutoff: float = 85.0, embed_cutoff: float = 0.78) -> Optional[str]:
    if not token: return None
    tok = token.strip().lower()
    if not tok: return None
    if tok in SKILL_ALIASES: return SKILL_ALIASES[tok]
    if tok in CANONICAL_SKILLS: return tok

    if CANONICAL_LIST:
        if RAPIDFUZZ: 
            try:
                match, score, _ = rf_process.extractOne(tok, CANONICAL_LIST, scorer=rf_fuzz.token_sort_ratio)
                if match and score >= fuzzy_cutoff:
                    return match
            except Exception:
                pass
        else:
            try:
                matches = get_close_matches(tok, CANONICAL_LIST, n=1, cutoff=(fuzzy_cutoff/100.0))
                if matches:
                    return matches[0]
            except Exception:
                pass

    global _skill_emb_cache
    if embedding_model is not None and CANONICAL_LIST:
        try:
            if _skill_emb_cache is None:
                _skill_emb_cache = {s: _np.array(embedding_model.encode(s), dtype=_np.float32) for s in CANONICAL_LIST}
            v_tok = _np.array(embedding_model.encode(tok), dtype=_np.float32)
            best_sim = 0.0; best_skill = None
            for s, vec in _skill_emb_cache.items():
                denom = (_np.linalg.norm(v_tok) * _np.linalg.norm(vec))
                if denom == 0: continue
                sim = float(_np.dot(v_tok, vec) / denom)
                if sim > best_sim:
                    best_sim = sim; best_skill = s
            if best_sim >= embed_cutoff:
                return best_skill
        except Exception:
            pass
    return None

# ---------------- Section splitting & headings ----------------
def split_into_sections(text: str):
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    if not lines:
        return [("body", text or "")]
    heading_indices = []
    for i, ln in enumerate(lines):
        s = ln.strip()
        if not s: continue
        if (len(s) < 90) and (s.isupper() or re.match(r'^[A-Z][a-z]+( [A-Za-z0-9&/-]+)*$', s) or (':' in s and len(s.split(':')[0].split()) <= 6)):
            heading_indices.append((i, s))
    if not heading_indices:
        for i, ln in enumerate(lines):
            if ':' in ln and len(ln) < 120:
                heading_indices.append((i, ln.split(':',1)[0].strip()))
    if not heading_indices:
        return [("body", text or "")]
    sections = []
    for idx, (i, hdr) in enumerate(heading_indices):
        start = i + 1
        end = heading_indices[idx+1][0] if idx+1 < len(heading_indices) else len(lines)
        body = "\n".join(lines[start:end]).strip()
        sections.append((hdr.strip(), body))
    merged = []
    for hdr, body in sections:
        if merged and len(body.split()) < 3:
            prev_hdr, prev_body = merged[-1]
            merged[-1] = (prev_hdr, (prev_body + "\n" + hdr + "\n" + body).strip())
        else:
            merged.append((hdr, body))
    return merged

_HEADINGS_PLAIN = [
    "professional experience","work experience","experience","employment history","work history",
    "professional history","experience & qualifications","experience and qualifications",
    "work & experience","professional experience and qualifications"
]

def canonicalize_heading(h: str):
    if not h: return "body"
    hu = h.upper()
    if "SKILL" in hu: return "skills"
    if "EXPERIENCE" in hu or "EMPLOY" in hu: return "experience"
    if "EDUC" in hu: return "education"
    if "CERTIF" in hu or "LICENSE" in hu: return "certifications"
    if "PROJECT" in hu: return "projects"
    if "CONTACT" in hu or "EMAIL" in hu or "PHONE" in hu: return "contact"
    if "SUMMARY" in hu or "PROFILE" in hu or "OBJECTIVE" in hu: return "summary"
    try:
        h_low = h.strip().lower()
        close = get_close_matches(h_low, _HEADINGS_PLAIN, n=1, cutoff=0.75)
        if close:
            cand = close[0]
            if "experience" in cand: return "experience"
            if "skill" in cand: return "skills"
            if "education" in cand: return "education"
            if "project" in cand: return "projects"
            if "contact" in cand: return "contact"
            if "summary" in cand or "profile" in cand: return "summary"
    except Exception:
        pass
    return h.lower().strip()

def extract_skills_from_section(text: str) -> List[str]:
    """
    Conservative heuristic extractor with BETTER filtering.
    - Prefers explicit comma/pipe/slash lists
    - Filters out section headers and meta text
    - Only returns actual technical/soft skills
    """
    if not text:
        return []
    
    lines = [ln.strip(" \tâ€¢Â·-:") for ln in text.splitlines() if ln.strip()]
    skills = []
    
    # Known tech substrings to help identify real skills
    tech_whitelist_subs = (
        "js","py","python","sql","aws","azure","gcp","c#","csharp","c++","cpp","java","scala",
        "spark","hadoop","react","node","django","flask","linux","k8s","kubernetes","docker",
        "php","html","css","json","xml","git","jira","mysql","postgres","oracle","mongodb",
        "nosql","jenkins","ansible","terraform","tableau","powerbi","pyspark","dbt","databricks"
    )
    
    # Meta keywords to SKIP (not actual skills)
    skip_patterns = re.compile(
        r'^\s*(experience|years?|month|week|day|project|tool|platform|'
        r'framework|library|package|software|application|system|'
        r'certification|certificate|license|degree|award|certification|'
        r'languages?|technologies?|skills?|expertise|proficiency)\s*:?\s*$',
        re.IGNORECASE
    )
    
    for ln in lines:
        # Skip lines that are section headers or meta-text
        if skip_patterns.match(ln):
            continue
        
        # Skip very long lines (likely descriptions, not skill lists)
        if len(ln.split()) > 20:
            continue
        
        # Handle explicit lists (comma/semicolon separated)
        if (',' in ln or ';' in ln) and len(ln) < 400:
            parts = [p.strip() for p in re.split(r'[,;]+', ln) if p.strip()]
            for p in parts:
                if 2 <= len(p.split()) <= 6:  # Multi-word skills OK
                    skills.append(p.lower())
            continue
        
        # Handle pipe/bullet separated lists
        if any(sep in ln for sep in ['|', '/', 'Â·', 'â€¢']):
            parts = re.split(r'[|/Â·â€¢]', ln)
            for p in parts:
                p = p.strip()
                if p and 2 <= len(p.split()) <= 6:
                    skills.append(p.lower())
            continue
        
        # Multi-word phrases (prefer these)
        mw = re.findall(r'\b[A-Za-z0-9\+\-#\.]{2,}(?:[ \t]+[A-Za-z0-9\+\-#\.]{2,})+\b', ln)
        for m in mw:
            skills.append(m.strip().lower())
        
        # Single-word tokens ONLY if they look tech-like
        sw = re.findall(r'\b[a-zA-Z0-9\+\-#\.]{2,}\b', ln)
        for s in sw:
            s_l = s.strip().lower()
            # Include if: has digits OR contains known tech substring
            if re.search(r'\d', s_l) or any(x in s_l for x in tech_whitelist_subs):
                skills.append(s_l)
    
    # Clean and dedupe
    seen = set()
    out = []
    for s in skills:
        s2 = re.sub(r'^[\W_]+|[\W_]+$', '', s).strip()
        if not s2 or len(s2) <= 1:
            continue
        
        # Skip stopwords and non-skills
        if s2 in ("and", "or", "with", "experience", "years", "year", "skills", "skill",
                  "the", "a", "an", "in", "of", "for", "is", "are", "technologies",
                  "knowledge", "expertise", "proficiency"):
            continue
        
        if s2 not in seen:
            seen.add(s2)
            out.append(s2)
    
    return out

def extract_skills_from_section_combined(section_text: str, full_text: str = "") -> List[str]:
    """
    Combine heuristic extraction and optional model-based extraction.
    - Heuristic extraction is always performed.
    - If `skill_pipe` (model) is available, its outputs are merged (model first, heuristics second).
    - Results are post-processed via post_process_skills() and categorized with your CSV mapping.
    - IMPORTANT: mapping/categorization is called with fuzzy=False (no fuzzy/embedding mapping),
      so only exact/canonical skills from categorized.csv will be mapped into their categories.
    Returns final cleaned list (lower-cased canonical keys when possible + normalized tokens).
    """
    heur_skills = []
    model_skills = []

    source_text = (section_text or full_text or "").strip()
    try:
        heur_skills = extract_skills_from_section(section_text or full_text or "")
    except Exception as e:
        heur_skills = []
        print("Heuristic extraction error:", e)

    # run model-based extractor if available (optional)
    try:
        if skill_pipe is not None and source_text:
            # give the model a reasonably sized chunk
            chunk = source_text[:12000]
            ents = skill_pipe(chunk)
            for ent in ents:
                # model output shape may vary; attempt robust extraction
                w = None
                if isinstance(ent, dict):
                    w = ent.get("word") or ent.get("text") or ent.get("entity") or ent.get("label")
                elif isinstance(ent, (str,)):
                    w = ent
                if w:
                    w_str = re.sub(r'[\u0120\u2581]', ' ', str(w)).strip()
                    if w_str:
                        model_skills.append(w_str.lower())
    except Exception as e:
        # don't fail entire parsing when model fails
        print("skill_pipe run error:", e)
        model_skills = []

    # Merge: prefer model suggestions first (if any), then heuristics
    merged = []
    seen = set()
    for m in model_skills:
        mk = re.sub(r'^[\W_]+|[\W_]+$', '', m).strip().lower()
        if mk and mk not in seen:
            seen.add(mk); merged.append(mk)
    for h in heur_skills:
        hk = re.sub(r'^[\W_]+|[\W_]+$', '', h).strip().lower()
        if hk and hk not in seen:
            seen.add(hk); merged.append(hk)

    # Post-process tokens (cleaning / canonical mapping attempt inside post_process)
    try:
        cleaned = post_process_skills(merged, full_text=source_text, whitelist=None, fuzzy_cutoff=0.90)
    except Exception as e:
        print("post_process_skills failed:", e)
        # fallback: use merged list
        cleaned = merged

    # Finally, categorize using your CSV-based categorizer but without fuzzy mapping
    try:
        categorized = categorize_amjad_skills(cleaned, fuzzy=False)
        # you requested not to create new files here: we return cleaned canonical tokens
        # The UI code elsewhere will consume `categorized` if needed; for this function return cleaned list
    except Exception as e:
        # if categorization fails, keep cleaned tokens
        print("categorize_amjad_skills error:", e)
        categorized = None

    # Return cleaned final list (prefer canonical-mapped tokens if post_process_skills returned them)
    return cleaned

# --- Categorization helpers (condensed) ---------------------------------
def categorize_amjad_skills(amjad_skills, fuzzy=False, fuzz_cutoff=80):
    canon_set = set([s for s in CANONICAL_LIST]) if CANONICAL_LIST else set()
    aliases = {k.lower(): v.lower() for k, v in SKILL_ALIASES.items()} if SKILL_ALIASES else {}
    canon_to_cat = {k.lower(): v.lower() for k, v in SKILL_TO_CATEGORY.items()} if SKILL_TO_CATEGORY else {}
    result = {"technical": {k: [] for k in [
        "programming_languages","cloud","databases","data_engineering","devops_ci_cd","web_frameworks",
        "analytics_and_viz","testing_and_quality","storage_and_format","other_tech"]},
        "soft": [], "other": []}
    unmatched = []
    canon_keys = sorted(list(canon_set))
    for raw in amjad_skills:
        orig = str(raw).strip()
        if not orig: continue
        tok = orig.lower()
        matched_canon = None
        if tok in canon_set:
            matched_canon = tok
        elif tok in aliases:
            matched_canon = aliases[tok]
        else:
            for c in canon_keys:
                if c and (c in tok or tok in c):
                    matched_canon = c; break
        if not matched_canon and fuzzy and canon_keys:
            if RAPIDFUZZ:
                matches = rf_process.extract(tok, canon_keys, scorer=rf_fuzz.ratio, score_cutoff=fuzz_cutoff, limit=3)
                if matches: matched_canon = matches[0][0]
            else:
                close = get_close_matches(tok, canon_keys, n=1, cutoff=fuzz_cutoff/100.0)
                if close: matched_canon = close[0]
        if matched_canon:
            cat = canon_to_cat.get(matched_canon, "")
            mapped = cat if cat else None
            if mapped == "soft":
                result["soft"].append(orig)
            elif mapped in result["technical"]:
                result["technical"][mapped].append(orig)
            else:
                result["technical"]["other_tech"].append(orig)
        else:
            unmatched.append(orig)
            result["other"].append(orig)
    for k in list(result["technical"].keys()):
        result["technical"][k] = sorted(dict.fromkeys(result["technical"][k]))
    result["soft"] = sorted(dict.fromkeys(result["soft"]))
    result["other"] = sorted(dict.fromkeys(result["other"]))
    result["_meta"] = {"unmatched_count": len(unmatched), "unmatched_samples": unmatched[:40]}
    return result

def print_categorized(result):
    print("Categorized skills:")
    for bucket, skills in result["technical"].items():
        if skills:
            print(f" {bucket}:")
            for sk in skills:
                print(" -", sk)
    if result["soft"]:
        print("Soft:")
        for sk in result["soft"]:
            print(" -", sk)
    if result["other"]:
        print("Other (not in CSV):")
        for sk in result["other"]:
            print(" -", sk)

def post_process_skills(raw_list: List[str], full_text: str = "", whitelist: set = None, fuzzy_cutoff: float = 0.78) -> List[str]:
    if not raw_list: return []
    out = []
    email_re = re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-z]{2,}', re.I)
    phone_re = re.compile(r'(\+?\d[\d\-\s()]{5,}\d)')
    seen = set()
    for tok in raw_list:
        if not tok: continue
        t = tok.strip()
        if email_re.search(t) or phone_re.search(t): continue
        if re.fullmatch(r'\d{4}', t) or re.fullmatch(r'[\d\-\./]{2,}', t): continue

        if t.lower().startswith("and ") or " and " in t.lower():
            parts = [p.strip() for p in re.split(r'\band\b', t, flags=re.I) if p.strip()]
            for p in parts:
                mapped = map_token_to_canonical(p, fuzzy_cutoff=85.0)
                if mapped and mapped not in seen:
                    out.append(mapped); seen.add(mapped)
            continue

        if len(t.split()) > 6:
            techs = re.findall(r'\b(java|python|scala|spark|hadoop|sql|pyspark|aws|azure|gcp|bigquery|jenkins|docker|kubernetes|dbt|databricks|snowflake|tableau)\b', t, flags=re.I)
            if techs:
                for tt in techs:
                    mapped = map_token_to_canonical(tt, fuzzy_cutoff=90.0)
                    if mapped and mapped not in seen:
                        out.append(mapped); seen.add(mapped)
                continue
            else:
                continue

        tl = t.lower()
        if tl in ("and","with","experience","skills","skill","years","year","the","a","an","in","on","of","for","is","are"):
            continue

        mapped = map_token_to_canonical(tl, fuzzy_cutoff=85.0)
        if mapped:
            if mapped not in seen:
                out.append(mapped); seen.add(mapped)
            continue

        if any(x in tl for x in ("python","java","scala","spark","hadoop","sql","aws","azure","gcp","docker","kubernetes","jenkins","react","node","django","flask")):
            norm = re.sub(r'^[^a-z0-9]+|[^a-z0-9]+$', '', tl)
            if norm and norm not in seen:
                out.append(norm); seen.add(norm)
            continue
    return out

# ---------------- Text extraction ----------------
logging.basicConfig(level=logging.INFO)
log = logging.getLogger(__name__)

DATE_RANGE_RE = re.compile(
    r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4}|\d{1,2}/\d{4})\s*(?:[-–—to]{1,4})\s*(?:Present|present|Current|current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4}|\d{1,2}/\d{4})',
    flags=re.I
)
SINGLE_DATE_RE = re.compile(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4}|\d{1,2}/\d{4}', flags=re.I)
_BULLET_RE = re.compile(r'^[\u2022\-\*\•\u25D8\u00B0\u00B70-9\)]\s+')

def _is_probably_binary(s: bytes) -> bool:
    if not s: return False
    if b'PK\x03\x04' in s[:4096]: return True
    total = len(s)
    nonprint = sum(1 for c in s if c < 32 and c not in (9,10,13))
    if total > 0 and (nonprint / float(total)) > 0.15:
        return True
    return False

def _clean_text_for_output(text: str) -> str:
    if not text: return ""
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[\t\f\v]+', ' ', text)
    text = ''.join(ch if (31 < ord(ch) < 127 or ord(ch) == 10) else ' ' for ch in text)
    text = re.sub(r'\n\s+\n', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()

def safe_text_extract(path: str) -> str:
    text = ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            try:
                import docx2txt
                text = docx2txt.process(path) or ""
                text = _clean_text_for_output(text)
                if text: return text
            except Exception:
                pass
            try:
                from docx import Document
                doc = Document(path)
                paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text: paragraphs.append(row_text)
                text = "\n".join(paragraphs)
                text = _clean_text_for_output(text)
                if text: return text
            except Exception:
                pass
            try:
                import zipfile, xml.etree.ElementTree as ET
                with zipfile.ZipFile(path, 'r') as z:
                    if "word/document.xml" in z.namelist():
                        raw = z.read("word/document.xml")
                        raw = raw.decode('utf-8', errors='ignore') if isinstance(raw, bytes) else str(raw)
                        try:
                            root = ET.fromstring(raw)
                            texts = []
                            for node in root.iter():
                                tag = getattr(node, 'tag', '')
                                if tag.endswith('}t'):
                                    if node.text: texts.append(node.text)
                            text = "\n".join(t.strip() for t in texts if t and t.strip())
                            text = _clean_text_for_output(text)
                            if text: return text
                        except Exception:
                            plain = re.sub(r'<[^>]+>', ' ', raw)
                            plain = _clean_text_for_output(plain)
                            if plain: return plain
            except Exception:
                pass
            return ""

        if ext == ".pdf":
            try:
                import pdfplumber
                pages = []
                with pdfplumber.open(path) as pdf:
                    for p in pdf.pages:
                        try:
                            pt = p.extract_text() or ""
                        except Exception:
                            pt = ""
                        if pt and pt.strip():
                            pages.append(pt)
                text = "\n\n".join(pages).strip()
                text = _clean_text_for_output(text)
                if text: return text
            except Exception:
                pass
            try:
                from pdfminer.high_level import extract_text
                text = extract_text(path) or ""
                text = _clean_text_for_output(text)
                if text: return text
            except Exception:
                pass
            return ""

        # Plain text or other
        try:
            with open(path, 'rb') as fh:
                raw = fh.read()
            if _is_probably_binary(raw):
                raise ValueError("File looks binary; try specialized extractor")
            try:
                candidate = raw.decode('utf-8')
            except Exception:
                try:
                    candidate = raw.decode('latin-1')
                except Exception:
                    candidate = ""
            text = _clean_text_for_output(candidate)
            if text: return text
        except Exception:
            pass

        try:
            import textract
            raw = textract.process(path)
            if isinstance(raw, bytes):
                raw = raw.decode('utf-8', errors='ignore')
            text = _clean_text_for_output(str(raw))
            if text: return text
        except Exception:
            pass
    except Exception as e:
        log.debug("safe_text_extract error for %s: %s", path, e)
    return ""

# ---------------- Candidate/name/contact/location helpers ----------------
def extract_locations_from_text(text: str):
    try:
        doc = nlp(text or "")
        locs = set()
        for ent in getattr(doc, "ents", []):
            if getattr(ent, "label_", "") in ("GPE", "LOC"):
                locs.add(ent.text.strip())
        return [l.lower() for l in sorted(locs)] if locs else []
    except Exception:
        return []

def extract_candidate_name_from_text(full_text: str, filename: str = None):
    """
    Improved name extraction that skips section headings.
    Checks first 15 lines for a line that looks like a name (not a heading).
    """
    lines = [ln.strip() for ln in (full_text or "").splitlines() if ln.strip()]
    skip_kw = re.compile(
        r'(?i)\b(resume|cv|curriculum vitae|profile|project|objective|summary|contact|'
        r'phone|email|address|linkedin|github|technical|manager|engineer|developer|'
        r'architect|specialist|lead|senior|junior|associate|coordinator|analyst)\b'
    )
    
    # Check first 15 lines for name-like content
    for ln in lines[:15]:
        # Skip if line contains keywords suggesting it's a heading/title
        if skip_kw.search(ln):
            continue
        
        words = ln.split()
        
        # Name heuristic: 1-4 words, all capitalized, no digits, length < 80
        if (1 <= len(words) <= 4 and 
            len(ln) < 80 and
            not any(ch.isdigit() for ch in ln)):
            
            # All words start with capital (but allow "van", "de", etc.)
            if all((w and w[0].isupper()) for w in words):
                return " ".join([w.capitalize() for w in ln.split()])
    
    # Fallback: try spaCy NER if available
    try:
        doc = nlp("\n".join(lines[:30]))
        persons = [ent.text.strip() for ent in getattr(doc, "ents", []) 
                  if getattr(ent, "label_", "") == "PERSON"]
        for p in sorted(persons, key=lambda s: -len(s)):
            if not skip_kw.search(p):
                return p
    except Exception:
        pass
    
    # Fallback: clean filename
    if filename:
        base = os.path.splitext(os.path.basename(filename))[0]
        base_clean = re.sub(r'[_\-.]+', ' ', base)
        base_clean = re.sub(r'(?i)\b(resume|cv|final|de|profile|pdf|docx)\b', '', base_clean)
        if base_clean.strip():
            return " ".join([w.capitalize() for w in base_clean.split()])
    
    return "Unknown"


try:
    from dateutil import parser as _dateutil_parser
    _HAS_DATEUTIL = True
except Exception:
    _HAS_DATEUTIL = False

# helper: parse a single token to a date (best-effort)
def _try_parse_date_token(token: str):
    token = (token or "").strip().strip(".,;()[]")
    if not token:
        return None
    if token.lower() in ("present", "current", "to date", "now"):
        return "PRESENT"
    # try dateutil first
    if _HAS_DATEUTIL:
        try:
            dt = _dateutil_parser.parse(token, default=datetime.datetime(1900,1,1), fuzzy=True)
            return dt.date()
        except Exception:
            pass
    # mm/yyyy or m/yyyy
    m = _re.search(r'(\d{1,2})[\/\-](\d{4})', token)
    if m:
        try:
            mon = int(m.group(1)); year = int(m.group(2))
            return datetime.date(year, max(1, min(12, mon)), 1)
        except Exception:
            pass
    # year only
    m2 = _re.search(r'(\d{4})', token)
    if m2:
        try:
            return datetime.date(int(m2.group(1)), 1, 1)
        except Exception:
            pass
    return None
def _merge_intervals_and_total_years(intervals):
    """
    Merge overlapping intervals and compute total years.
    intervals: list of (start_date, end_date) tuples
    returns: float years (1 decimal)
    """
    if not intervals:
        return 0.0
    
    clean = [(s, e) for (s, e) in intervals if s and e and e >= s]
    if not clean:
        return 0.0
    
    clean.sort(key=lambda x: x[0])
    
    # Merge overlapping intervals
    merged = []
    cur_s, cur_e = clean[0]
    for s, e in clean[1:]:
        if s <= cur_e:
            cur_e = max(cur_e, e)
        else:
            merged.append((cur_s, cur_e))
            cur_s, cur_e = s, e
    merged.append((cur_s, cur_e))
    
    # Sum total days
    total_days = sum((e - s).days for s, e in merged)
    years = round(total_days / 365.0, 1)
    
    # Sanity checks
    if years < 0:
        years = 0.0
    elif years > 60:
        years = 60.0
    
    return float(years)

def extract_experience_from_section(text: str):
    """
    Improved experience parsing with robust date handling and duration calculation.
    Returns list of experience entries with calculated duration_years.
    """
    if not text or not isinstance(text, str):
        return []

    paras = [p.strip() for p in _re.split(r'\n{2,}', text) if p.strip()]
    out = []
    
    # More flexible date patterns
    RANGE_RE = _re.compile(
        r'(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|'
        r'\d{1,2}/\d{4}|\d{4})\s*'
        r'(?:[-â€"to\s]+|–)\s*'
        r'(?P<end>Present|present|Current|current|Now|now|'
        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|'
        r'\d{1,2}/\d{4}|\d{4})',
        flags=_re.IGNORECASE
    )
    
    # Explicit duration mention
    YEARS_RE = _re.compile(r'(\d+(?:\.\d+)?)\s+years?', flags=_re.IGNORECASE)
    
    MONTH_MAP = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
    }

    def parse_date(date_str):
        """Parse a date string to datetime.date object."""
        date_str = (date_str or "").strip().lower()
        if not date_str:
            return None
        
        if date_str in ('present', 'current', 'now'):
            return datetime.date.today()
        
        # "Jan 2020" format
        m = _re.match(r'([a-z]{3})[a-z]*\.?\s*(\d{4})', date_str)
        if m:
            month_str, year = m.groups()
            month = MONTH_MAP.get(month_str[:3], 1)
            try:
                return datetime.date(int(year), month, 1)
            except:
                return None
        
        # "01/2020" format
        m = _re.match(r'(\d{1,2})/(\d{4})', date_str)
        if m:
            month, year = m.groups()
            try:
                return datetime.date(int(year), int(month), 1)
            except:
                return None
        
        # Year only "2020"
        m = _re.match(r'(\d{4})', date_str)
        if m:
            try:
                return datetime.date(int(m.group(1)), 1, 1)
            except:
                return None
        
        return None

    def calculate_duration(start_date, end_date):
        """Calculate duration in years between two dates."""
        if not start_date or not end_date or end_date < start_date:
            return None
        
        delta = end_date - start_date
        years = delta.days / 365.25
        return round(max(0.1, years), 1)  # Minimum 0.1 years

    for p in paras:
        entry = {
            "raw": p,
            "title": "",
            "company": "",
            "start_date": None,
            "end_date": None,
            "duration_years": None,
            "location": "",
            "summary": p
        }
        
        # Try to find date range
        m = RANGE_RE.search(p)
        if m:
            start_str = m.group("start")
            end_str = m.group("end")
            
            start_date = parse_date(start_str)
            end_date = parse_date(end_str)
            
            if start_date:
                entry["start_date"] = start_date.isoformat()
            if end_date:
                entry["end_date"] = end_date.isoformat()
            
            # Calculate duration
            if start_date and end_date:
                duration = calculate_duration(start_date, end_date)
                if duration:
                    entry["duration_years"] = duration
        
        # Fallback: look for explicit "X years" mention
        if entry["duration_years"] is None:
            my = YEARS_RE.search(p)
            if my:
                try:
                    years = float(my.group(1))
                    entry["duration_years"] = round(min(60.0, max(0.1, years)), 1)
                except:
                    pass
        
        # Extract title and company from first line
        first_line = (p.splitlines() or [""])[0].strip()
        
        if ' - ' in first_line or ' â€" ' in first_line:
            parts = _re.split(r'\s*[-â€"]\s*', first_line, maxsplit=1)
            if len(parts) == 2:
                left, right = parts[0].strip(), parts[1].strip()
                if any(x in left.lower() for x in ('inc', 'llc', 'ltd', 'corp', 'company', 'group', 'systems', 'services')):
                    entry["company"] = left
                    entry["title"] = right
                else:
                    entry["title"] = left
                    entry["company"] = right
        elif '|' in first_line:
            parts = first_line.split('|', maxsplit=1)
            entry["title"] = parts[0].strip()
            entry["company"] = parts[1].strip() if len(parts) > 1 else ""
        else:
            entry["title"] = first_line
        
        out.append(entry)
    
    return out
def categorize_skills_for_resume(skills: List[str], full_text: str = "") -> Dict[str, Any]:
    technical_out: Dict[str, List[str]] = {k: [] for k in [
        "programming_languages","cloud","databases","data_engineering","devops_ci_cd",
        "web_frameworks","analytics_and_viz","testing_and_quality","storage_and_format","other_tech"
    ]}
    soft_out: List[str] = []
    other_out: List[str] = []
    for tok in skills:
        t = tok.lower().strip()
        if t in _SOFT_SKILL_KEYWORDS:
            soft_out.append(tok); continue
        mapped = _FLAT_SKILL_TO_CATEGORY.get(t)
        if mapped and mapped in technical_out:
            technical_out[mapped].append(tok); continue
        if any(x in t for x in ("communication","leadership","teamwork","mentoring","presentation","collaboration")):
            soft_out.append(tok); continue
        substr_to_bucket = {
            "python": "programming_languages","java": "programming_languages","c++": "programming_languages",
            "sql": "databases","mysql": "databases","postgres": "databases","mongodb": "databases",
            "aws": "cloud","azure": "cloud","gcp": "cloud","spark": "data_engineering","hadoop": "data_engineering",
            "airflow": "data_engineering","docker": "devops_ci_cd","kubernetes": "devops_ci_cd","jenkins": "devops_ci_cd",
            "react": "web_frameworks","angular": "web_frameworks","django": "web_frameworks","flask": "web_frameworks",
            "tableau": "analytics_and_viz","powerbi": "analytics_and_viz","pytest": "testing_and_quality",
            "parquet": "storage_and_format","json": "storage_and_format","csv": "storage_and_format"
        }
        bucket = None
        for sub, b in substr_to_bucket.items():
            if sub in t:
                bucket = b; break
        if bucket:
            technical_out[bucket].append(tok)
        else:
            if 1 <= len(t.split()) <= 4:
                other_out.append(tok)
    for k in list(technical_out.keys()):
        technical_out[k] = sorted(dict.fromkeys(technical_out[k]))
    soft_out = sorted(dict.fromkeys(soft_out))
    other_out = sorted(dict.fromkeys(other_out))
    return {"technical": technical_out, "soft": soft_out, "other": other_out}

def format_candidate_json(parsed_payload: Dict[str, Any]) -> str:
    """
    Convert a parsed candidate payload into a JSON string suitable for UI display.
    """
    name = parsed_payload.get("candidate_name") or parsed_payload.get("candidate_id") or "Unknown"
    email = parsed_payload.get("email")
    phone = parsed_payload.get("phone")
    linkedin = parsed_payload.get("linkedin") or parsed_payload.get("profile_link") or None

    # normalize locations -> list of non-empty strings
    locations_raw = parsed_payload.get("locations") or []
    if isinstance(locations_raw, str):
        locations = [locations_raw]
    elif isinstance(locations_raw, (list, tuple)):
        locations = [str(l).strip() for l in locations_raw if l and str(l).strip()]
    else:
        locations = []

    # safe summary
    summary = parsed_payload.get("summary") or parsed_payload.get("objective") or ""
    try:
        summary = str(summary)[:2000]
    except Exception:
        summary = ""

    # skills: prefer raw list for categorization, but include normalized skills_set if available
    skills_list = parsed_payload.get("skills") or []
    # ensure skills_list is a list of strings
    if isinstance(skills_list, (str,)):
        skills_list = [skills_list]
    skills_list = [str(s).strip() for s in (skills_list or []) if s]

    # categorized skills (safe call)
    try:
        categorized = categorize_skills_for_resume(skills_list, full_text=parsed_payload.get("full_text", ""))
    except Exception:
        # fallback: minimal structure
        categorized = {"technical": skills_list, "soft": [], "other": []}

    # include normalized skills_set if present
    skills_set = parsed_payload.get("skills_set")
    if isinstance(skills_set, (list, tuple)):
        skills_set = [str(s).strip().lower() for s in skills_set if s]
    else:
        skills_set = [s.lower() for s in skills_list]

    # experience: return as-is but protect against non-list shapes
    experience = parsed_payload.get("experience") or []
    if not isinstance(experience, (list, tuple)):
        experience = [experience]

    # total years (prefer canonical field if set)
    total_years = parsed_payload.get("total_years_experience")
    try:
        total_years = float(total_years) if total_years is not None else 0.0
        if total_years < 0 or total_years > 60:
            total_years = round(max(0.0, min(60.0, total_years)), 1)
    except Exception:
        total_years = 0.0

    out = {
        "candidate_name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "locations": locations,
        "summary": summary,
        "skills": categorized,
        "skills_set": skills_set,
        "total_years_experience": total_years,
        "experience": experience,
    }

    return json.dumps(out, indent=2, ensure_ascii=False)
def compute_query_skill_similarity(query: str, candidate_skills: List[str], embedding_model) -> float:
    """
    Compute semantic similarity between query and candidate's skill set.
    Uses embeddings to compare query against skills dynamically.
    
    Returns: similarity score [0.0, 1.0]
    """
    if not candidate_skills or not query:
        return 0.0
    
    try:
        # Embed query once
        query_vec = _np.array(embedding_model.encode(query), dtype=_np.float32)
        query_vec = query_vec / (_np.linalg.norm(query_vec) + 1e-9)
        
        # Embed all candidate skills (batch for efficiency)
        valid_skills = [s for s in candidate_skills if s and len(str(s).strip()) > 1]
        if not valid_skills:
            return 0.0
            
        skill_embeddings = embedding_model.encode(valid_skills[:50])  # limit to top 50 skills
        if not isinstance(skill_embeddings, _np.ndarray):
            skill_embeddings = _np.array(skill_embeddings)
        
        # Normalize skill vectors
        norms = _np.linalg.norm(skill_embeddings, axis=1, keepdims=True)
        norms = _np.where(norms == 0, 1.0, norms)
        skill_embeddings = skill_embeddings / norms
        
        # Compute cosine similarities
        similarities = _np.dot(skill_embeddings, query_vec)
        
        # Use top-k average similarity (more robust than max)
        top_k = min(5, len(similarities))
        top_sims = _np.partition(similarities, -top_k)[-top_k:]
        avg_sim = float(_np.mean(top_sims))
        
        return max(0.0, min(1.0, avg_sim))  # clamp to [0,1]
        
    except Exception as e:
        return 0.0


def compute_role_relevance(query: str, candidate_text: str, embedding_model) -> float:
    """
    Compute semantic relevance between query and candidate's role/title text.
    Uses embeddings to determine if candidate matches the role dynamically.
    
    Returns: relevance score [0.0, 1.0]
    """
    if not query or not candidate_text:
        return 0.0
    
    try:
        # Embed both texts
        query_vec = _np.array(embedding_model.encode(query), dtype=_np.float32)
        text_vec = _np.array(embedding_model.encode(candidate_text[:500]), dtype=_np.float32)
        
        # Normalize
        query_vec = query_vec / (_np.linalg.norm(query_vec) + 1e-9)
        text_vec = text_vec / (_np.linalg.norm(text_vec) + 1e-9)
        
        # Cosine similarity
        sim = float(_np.dot(query_vec, text_vec))
        return max(0.0, min(1.0, sim))
        
    except Exception as e:
        return 0.0
def compute_semantic_relevance(query: str, candidate_text: str, embedding_model) -> float:
    """
    Direct semantic relevance between query and candidate.
    No templates, no keywords - pure embedding similarity.
    
    Returns: relevance score [0.0, 1.0]
    """
    if not query or not candidate_text:
        return 0.0
    
    try:
        # Embed both query and candidate text
        query_vec = _np.array(embedding_model.encode(query.strip()), dtype=_np.float32)
        
        # Use a larger chunk of candidate text for better context
        candidate_vec = _np.array(
            embedding_model.encode(candidate_text[:3000].strip()), 
            dtype=_np.float32
        )
        
        # Normalize vectors
        query_vec = query_vec / (_np.linalg.norm(query_vec) + 1e-9)
        candidate_vec = candidate_vec / (_np.linalg.norm(candidate_vec) + 1e-9)
        
        # Compute cosine similarity
        similarity = float(_np.dot(query_vec, candidate_vec))
        
        return max(0.0, min(1.0, similarity))
        
    except Exception as e:
        return 0.0


def compute_section_relevance(query: str, sections: Dict[str, str], embedding_model) -> Dict[str, float]:
    """
    Compute relevance for each resume section.
    Helps identify which parts of resume match the query best.
    """
    section_scores = {}
    
    if not query or not sections:
        return section_scores
    
    try:
        query_vec = _np.array(embedding_model.encode(query.strip()), dtype=_np.float32)
        query_vec = query_vec / (_np.linalg.norm(query_vec) + 1e-9)
        
        for section_name, section_text in sections.items():
            if not section_text or len(str(section_text).strip()) < 20:
                continue
            
            try:
                section_vec = _np.array(
                    embedding_model.encode(str(section_text)[:1000]), 
                    dtype=_np.float32
                )
                section_vec = section_vec / (_np.linalg.norm(section_vec) + 1e-9)
                
                similarity = float(_np.dot(query_vec, section_vec))
                section_scores[section_name] = max(0.0, min(1.0, similarity))
            except:
                continue
        
        return section_scores
        
    except Exception:
        return section_scores


def semantic_search(query: str,
                    top_k: int = 20,
                    qdrant_limit: int = 500,
                    debug: bool = False,
                    filter_candidate_name: Optional[str] = None,
                    filter_location: Optional[str] = None,
                    filter_skill: Optional[str] = None,
                    min_years_experience: Optional[float] = None,
                    relevance_threshold: float = 0.30) -> List[Dict[str, Any]]:
    """
    Pure semantic search - with better fallback handling.
    """
    import uuid
    from collections import defaultdict
    
    global qdrant_client, embedding_model, QDRANT_COLLECTION
    
    if qdrant_client is None or embedding_model is None:
        if debug:
            print("[semantic_search] qdrant_client or embedding_model not initialized.")
        return []
    
    if debug:
        print(f"[semantic_search] Query: '{query}', Threshold: {relevance_threshold}")
    
    # 1) Embed query
    try:
        qv = embedding_model.encode(query)
        qv = _np.asarray(qv, dtype=float).flatten()
        nq = _np.linalg.norm(qv)
        if nq > 0:
            qv = (qv / (nq + 1e-12)).tolist()
        else:
            qv = qv.tolist()
    except Exception as e:
        if debug:
            print("[semantic_search] embedding failed:", e)
        return []
    
    # 2) Call Qdrant
    resp = None
    try:
        resp = qdrant_client.search(
            collection_name=QDRANT_COLLECTION,
            query_vector=qv,
            limit=max(qdrant_limit, top_k * 15),
            with_payload=True
        )
    except Exception as e:
        if debug:
            print("[semantic_search] Qdrant search failed:", e)
        return []
    
    if not resp:
        if debug:
            print("[semantic_search] no hits from qdrant.")
        return []
    
    # 3) Normalize hits
    hits = []
    for r in resp:
        payload = {}
        score = None
        pid = None
        if hasattr(r, "payload"):
            payload = r.payload or {}
            score = getattr(r, "score", None)
            pid = getattr(r, "id", None)
        elif isinstance(r, dict):
            payload = r.get("payload") or {}
            score = r.get("score")
            pid = r.get("id")
        
        try:
            numeric_score = float(score) if score is not None else 0.0
        except Exception:
            numeric_score = 0.0
        
        hits.append({"id": pid, "score": numeric_score, "payload": payload})
    
    if debug:
        print(f"[semantic_search] Got {len(hits)} raw hits from Qdrant")
    
    # 4) Aggregate by candidate
    def aggregate_candidates(hits_list, query_text, threshold):
        scores_by_candidate = defaultdict(list)
        payload_by_candidate = {}
        
        for h in hits_list:
            pl = h.get("payload") or {}
            cid = (pl.get("candidate_id") or pl.get("email") or 
                   pl.get("candidate_name") or h.get("id") or str(uuid.uuid4()))
            cid = str(cid).strip().lower()
            s = float(h.get("score") or 0.0)
            scores_by_candidate[cid].append(s)
            
            prev = payload_by_candidate.get(cid)
            if prev is None or s > max(scores_by_candidate[cid][:-1], default=0):
                payload_by_candidate[cid] = pl
        
        aggregated = []
        
        for cid, sc_list in scores_by_candidate.items():
            pl = payload_by_candidate[cid] or {}
            
            # Get candidate context
            full_text = pl.get("full_text") or ""
            sections = pl.get("sections") or {}
            skills = pl.get("skills_set") or pl.get("skills") or []
            
            # Build candidate text
            candidate_context_parts = []
            
            if sections.get("summary"):
                candidate_context_parts.append(str(sections["summary"])[:500])
            if sections.get("experience"):
                candidate_context_parts.append(str(sections["experience"])[:1500])
            
            if isinstance(skills, list):
                skills_text = " ".join([str(s) for s in skills[:50] if s])
                candidate_context_parts.append(f"Skills: {skills_text}")
            
            if not candidate_context_parts and full_text:
                candidate_context_parts.append(full_text[:2000])
            
            candidate_text = " ".join(candidate_context_parts)
            
            # Compute semantic relevance (with fallback)
            try:
                semantic_score = compute_semantic_relevance(
                    query_text, 
                    candidate_text, 
                    embedding_model
                )
            except Exception as e:
                if debug:
                    print(f"[semantic] compute_semantic_relevance failed for {pl.get('candidate_name')}: {e}")
                # FALLBACK: use vector similarity instead of failing
                semantic_score = max(sc_list) if sc_list else 0.0
            
            # Don't filter yet - compute full score first
            sc_sorted = sorted(sc_list, reverse=True)
            top_chunk_scores = sc_sorted[:3]
            base_score = sum(top_chunk_scores) / len(top_chunk_scores)
            
            # Section relevance
            try:
                section_scores = compute_section_relevance(query_text, sections, embedding_model)
                relevant_sections = [s for s, score in section_scores.items() if score > 0.35]
                section_boost = min(0.10, len(relevant_sections) * 0.025)
            except Exception:
                section_scores = {}
                relevant_sections = []
                section_boost = 0.0
            
            # Skill match
            try:
                skill_similarity = compute_query_skill_similarity(
                    query_text, 
                    skills if isinstance(skills, list) else [], 
                    embedding_model
                )
                skill_boost = skill_similarity * 0.12
            except Exception:
                skill_boost = 0.0
            
            # Experience boost
            years = pl.get("total_years_experience", 0.0)
            try:
                years_float = float(years) if years else 0.0
                exp_boost = min(0.05, 0.015 * _np.log1p(max(0, years_float - 1)))
            except Exception:
                exp_boost = 0.0
            
            # FINAL SCORE
            final_score = (
                semantic_score * 0.70 +
                base_score * 0.15 +
                skill_boost +
                section_boost +
                exp_boost
            )
            
            # NOW filter by threshold
            if final_score < threshold:
                if debug:
                    print(f"[filtered] {pl.get('candidate_name', 'Unknown')}: {final_score:.3f} < {threshold}")
                continue
            
            aggregated.append({
                "candidate_id": cid,
                "score": float(final_score),
                "payload": pl,
                "semantic_relevance": float(semantic_score),
                "relevant_sections": relevant_sections,
                "section_scores": section_scores
            })
        
        return sorted(aggregated, key=lambda x: x["score"], reverse=True)
    
    ranked_candidates = aggregate_candidates(hits, query, relevance_threshold)
    
    if debug:
        print(f"[semantic_search] After aggregation: {len(ranked_candidates)} candidates")
        for i, c in enumerate(ranked_candidates[:5], 1):
            print(f"  {i}. {c['payload'].get('candidate_name', 'Unknown')}: "
                  f"score={c['score']:.3f}, semantic={c['semantic_relevance']:.3f}")
    
    # 5) Apply post-filters
    def _candidate_matches_filters(candidate_entry):
        pl = candidate_entry.get("payload") or {}
        
        if filter_skill:
            fs = str(filter_skill).strip().lower()
            skills_set = pl.get("skills_set") or pl.get("skills") or []
            if isinstance(skills_set, dict):
                skills_flat = []
                for v in skills_set.values():
                    if isinstance(v, (list, tuple)):
                        skills_flat.extend([str(x).lower() for x in v if x])
            elif isinstance(skills_set, (list, tuple)):
                skills_flat = [str(x).lower() for x in skills_set if x]
            else:
                skills_flat = []
            
            if not any(fs in s or s in fs for s in skills_flat):
                return False
        
        if filter_location:
            fl = str(filter_location).strip().lower()
            locs = pl.get("locations") or []
            if isinstance(locs, str):
                locs = [locs]
            locs_low = [str(x).lower() for x in locs if x]
            if not any(fl in l for l in locs_low):
                return False
        
        if min_years_experience is not None:
            try:
                required = float(min_years_experience)
                tys = pl.get("total_years_experience", 0.0)
                tys_f = float(tys or 0.0)
                if tys_f < required:
                    return False
            except Exception:
                pass
        
        return True
    
    filtered_candidates = [c for c in ranked_candidates if _candidate_matches_filters(c)]
    
    if debug:
        print(f"[semantic_search] After filters: {len(filtered_candidates)} candidates")
    
    # 6) Return top_k
    out = filtered_candidates[:top_k]
    return out

# ============================================================================
# Helper function for CLI to explain matches
# ============================================================================

def explain_match(candidate_result: Dict[str, Any], query: str) -> List[str]:
    """
    Generate natural language explanation for why a candidate matched.
    Uses the semantic relevance scores computed during search.
    """
    explanations = []
    
    semantic_score = candidate_result.get("semantic_relevance", 0.0)
    relevant_sections = candidate_result.get("relevant_sections", [])
    section_scores = candidate_result.get("section_scores", {})
    
    # Overall match quality
    if semantic_score > 0.45:
        explanations.append("• Strong semantic match with your requirements")
    elif semantic_score > 0.35:
        explanations.append("• Good alignment with your search criteria")
    else:
        explanations.append("• Moderate match with query intent")
    
    # Section-specific matches
    if relevant_sections:
        top_section = max(section_scores.items(), key=lambda x: x[1])[0]
        top_score = section_scores[top_section]
        if top_score > 0.40:
            section_name = top_section.replace("_", " ").title()
            explanations.append(f"• {section_name} section highly relevant ({top_score:.2f} similarity)")
    
    # Skills match
    payload = candidate_result.get("payload", {})
    skills = payload.get("skills_set") or payload.get("skills") or []
    if isinstance(skills, list) and len(skills) > 0:
        # Check for query terms in skills
        query_lower = query.lower()
        query_terms = set(query_lower.split())
        skill_matches = [s for s in skills if any(term in str(s).lower() for term in query_terms)]
        if skill_matches:
            explanations.append(f"• Relevant skills: {', '.join(str(s) for s in skill_matches[:4])}")
    
    # Experience
    years = payload.get("total_years_experience", 0)
    if years and years > 2:
        explanations.append(f"• {years:.1f} years of relevant experience")
    
    return explanations[:4] 


def parse_resume_file(path: str) -> dict:
    """
    Parse resume with FIXED experience calculation.
    """
    parsed = {
        "candidate_name": None,
        "candidate_id": os.path.basename(path),
        "email": None,
        "phone": None,
        "full_text": "",
        "locations": [],
        "sections": {},
        "skills": [],
        "skills_by_category": {"technical": {}, "soft": [], "other": []},
        "experience": [],
        "total_years_experience": 0.0
    }

    try:
        full_text = safe_text_extract(path) or ""
        parsed["full_text"] = full_text

        if full_text:
            m = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', full_text)
            parsed["email"] = m.group(0) if m else None
            m2 = re.search(r'(\+?\d[\d\-\.\s\(\)]{7,}\d)', full_text)
            parsed["phone"] = m2.group(0).strip() if m2 else None

        # Extract candidate name BEFORE parsing sections
        try:
            parsed_name = extract_candidate_name_from_text(parsed.get("full_text", ""), filename=path)
            if parsed_name and parsed_name != "Unknown":
                parsed["candidate_name"] = parsed_name
        except Exception:
            pass

        # Split into sections
        parsed_sections = {}
        if full_text:
            try:
                sections = split_into_sections(full_text)
                for hdr, body in sections:
                    key = canonicalize_heading(hdr) or "body"
                    parsed_sections[key] = parsed_sections.get(key, "") + ("\n" + body if parsed_sections.get(key) else body)
            except Exception:
                parsed_sections["body"] = full_text
        parsed["sections"] = parsed_sections

        # Extract skills (from skills section, with better filtering)
        try:
            skills_section_text = (parsed_sections.get("skills") or "").strip()
            if skills_section_text:
                extracted_skills = extract_skills_from_section_combined(skills_section_text, full_text=full_text)
            else:
                # Don't default to full_text for skill extraction - be conservative
                extracted_skills = extract_skills_from_section_combined("", full_text=full_text)

            cleaned = []
            seen = set()
            for s in (extracted_skills or []):
                st = str(s).strip()
                st = re.sub(r'^[\W_]+|[\W_]+$', '', st)
                if not st or len(st) < 2:
                    continue
                key = st.lower()
                if key not in seen:
                    seen.add(key)
                    cleaned.append(st)
            
            parsed["skills"] = cleaned
            parsed["skills_by_category"] = categorize_skills_for_resume(parsed["skills"], full_text=full_text)
        except Exception as e:
            log.debug("skill extraction failed for %s: %s", path, e)
            parsed["skills"] = []
            parsed["skills_by_category"] = {"technical": {}, "soft": [], "other": []}

        # Extract experience and calculate total years
        try:
            exp_text = parsed["sections"].get("experience", "").strip()
            if exp_text:
                parsed["experience"] = extract_experience_from_section(exp_text)
            else:
                parsed["experience"] = []

            # Calculate total years from structured experience entries
            intervals = []
            for e in parsed.get("experience", []):
                if isinstance(e, dict):
                    sd_token = e.get("start_date")
                    ed_token = e.get("end_date")

                    s_date = None
                    e_date = None

                    # Parse start date
                    if sd_token:
                        try:
                            if isinstance(sd_token, str):
                                s_date = datetime.date.fromisoformat(sd_token)
                            elif isinstance(sd_token, datetime.date):
                                s_date = sd_token
                        except Exception:
                            pass

                    # Parse end date
                    if ed_token:
                        try:
                            if isinstance(ed_token, str):
                                if ed_token.lower() in ("present", "current"):
                                    e_date = datetime.date.today()
                                else:
                                    e_date = datetime.date.fromisoformat(ed_token)
                            elif isinstance(ed_token, datetime.date):
                                e_date = ed_token
                        except Exception:
                            pass

                    # If start exists but no end, assume current
                    if s_date and not e_date:
                        e_date = datetime.date.today()

                    # Add valid interval
                    if s_date and e_date and e_date >= s_date:
                        intervals.append((s_date, e_date))

            parsed['total_years_experience'] = _merge_intervals_and_total_years(intervals)
        except Exception as e:
            log.debug("experience extraction failed: %s", e)
            parsed["experience"] = []
            parsed['total_years_experience'] = 0.0

        # Extract locations
        try:
            parsed["locations"] = extract_locations_from_text(full_text)
        except Exception:
            parsed["locations"] = []

    except Exception as e:
        log.debug("parse_resume_file error for %s: %s", path, e)

    return parsed